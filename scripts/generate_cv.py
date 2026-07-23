#!/usr/bin/env python3
"""Generate Shine Bo Bo's editable CV (DOCX) and website CV (PDF)."""

from pathlib import Path
from shutil import copy2
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUTPUT = ROOT / "output" / "docx" / "shine-bo-bo-cv.docx"
PDF_OUTPUT = ROOT / "output" / "pdf" / "shine-bo-bo-cv.pdf"
PUBLIC_OUTPUT = ROOT / "public" / "cv.pdf"

NAVY = "0F172A"
BLUE = "2563EB"
SLATE = "475569"
MUTED = "64748B"
LIGHT = "CBD5E1"

PROFILE = (
    "Computer Science graduate and third-year Digital Engineering student with project-based "
    "experience across full-stack web applications, backend APIs, AI/ML prototypes, IoT simulation, "
    "and practical PC hardware work. I combine technical learning with international student "
    "leadership, cross-cultural collaboration, and customer-facing problem-solving. I am seeking "
    "software development, full-stack/backend, or digital engineering internship and junior opportunities."
)

EDUCATION = [
    {
        "title": "Bachelor of Engineering (Digital Engineering), International Program",
        "meta": "Prince of Songkla University, Phuket Campus | College of Computing | 2024-Present | GPA 4.00",
        "bullets": [
            "Third-year student in a 120-credit, project- and real-case-based engineering program.",
            "Relevant study: computer programming, web development, UI/UX and design thinking, software engineering, computer architecture and operating systems, electronics and IoT, data communications, networking, cybersecurity, cryptography, project management, and quality assurance.",
            "Outstanding Academic Performance Award, academic year 2024, first-year Digital Engineering.",
        ],
    },
    {
        "title": "Bachelor of Science in Computer Science",
        "meta": "Myanmar Institute of Theology, Liberal Arts Program | Yangon, Myanmar | 2021-2025 | GPA 3.74",
        "bullets": [
            "Completed computing foundations in programming, software engineering, databases, operating systems, web and mobile development, and computer systems.",
            "Liberal-arts coursework: English, Korean Language Basics, Management, Introduction to Economics, Ethics, Human Rights, Gender Studies, Social Studies, Civic Education, Counselling, World Religion, Religion and Society, Peace Education, Wellbeing, Music Appreciation, Introduction to Christianity, Introduction to IT, and You and Your World.",
        ],
    },
]

SKILLS = [
    ("Programming", "TypeScript, JavaScript, Python, Java, PHP, C++, SQL, HTML, CSS"),
    ("Web interfaces", "React, Next.js, Tailwind CSS, responsive UI, Material UI, Framer Motion"),
    ("APIs and backend", "Node.js, Express, NestJS, FastAPI, Flask, Spring Boot, Laravel, REST APIs, authentication flows"),
    ("Data and AI/ML", "PostgreSQL, MySQL, MongoDB/Mongoose, Sequelize, SQLAlchemy, pandas, scikit-learn, TensorFlow/Keras, OpenCV, Streamlit"),
    ("IoT and workflow", "ESP32, MQTT, PlatformIO, Wokwi, Git, GitHub, Postman, npm, Vercel"),
    ("PC hardware", "Component selection, compatibility planning, custom PC assembly, diagnostics, troubleshooting, sourcing, resale, and basic after-sales support"),
    ("Coursework foundations", "Networking, cybersecurity, cryptography, computer architecture, operating systems, systems thinking, and engineering design"),
]

EXPERIENCE = [
    {
        "title": "President",
        "meta": "PSU Phuket International Student Club | Student Leadership | Current",
        "bullets": [
            "Lead an intercultural student community and help international students connect, participate, and feel represented in campus life.",
            "Coordinate club activities with students and university stakeholders while developing cross-cultural communication, planning, and team-leadership experience.",
        ],
    },
    {
        "title": "Research Intern - Flower Classification",
        "meta": "Myanmar Institute of Theology, Liberal Arts Program | Jan-Jul 2025",
        "bullets": [
            "Prepared and organized image data across five flower categories for a supervised academic AI/ML project.",
            "Trained and evaluated a MobileNetV2 transfer-learning workflow, documented results, and built a Streamlit prediction interface.",
        ],
    },
    {
        "title": "Freelance Gaming PC Buyer, Builder and Reseller",
        "meta": "Self-Employed / Family Business | Jan 2022-Jul 2024",
        "bullets": [
            "Matched components to customer needs, budgets, and compatibility constraints; sourced parts, assembled systems, and performed setup, diagnostics, and troubleshooting.",
            "Handled pricing, negotiation, customer communication, delivery, and basic after-sales support while assisting with day-to-day family-business operations.",
        ],
    },
    {
        "title": "Logistics and Delegate Servicing Team Member",
        "meta": "AIESEC in Myanmar | Volunteer Experience | Feb 2022-Jan 2023",
        "bullets": [
            "Supported youth-led events connected to the UN Sustainable Development Goals through logistics, delegate servicing, coordination, and participant-facing operations.",
            "Collaborated across functions and adapted to changing on-site needs to support smoother participant experiences.",
        ],
    },
]

PROJECTS = [
    {
        "title": "Real Estate Management System",
        "meta": "Team coursework | Next.js, TypeScript, NestJS, PostgreSQL, Sequelize, JWT | 2025",
        "text": "Contributed across frontend and backend repositories to buyer/dealer registration, role-based dashboards, authentication, property listings, and client/API integration.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/rems-full-stack-system",
    },
    {
        "title": "Cupidy",
        "meta": "External contributor | Python, FastAPI, SQLAlchemy, PostgreSQL, JWT | 2024",
        "text": "Contributed 25 GitHub-attributed commits covering signup/sign-in, profiles, photo upload, matching, likes, schemas, models, repository functions, and database setup.",
        "url": "https://github.com/Lin-Pyae/cupidy/commits/master/?author=ShineBo",
    },
    {
        "title": "Movie Recommendation Prototype",
        "meta": "External contributor | React, Flask, pandas, scikit-learn, Streamlit | 2024",
        "text": "Contributed eight public implementation commits spanning content-based recommendation experiments, Flask integration, React result/detail flows, and a Streamlit testing interface.",
        "url": "https://github.com/Lin-Pyae/Movie-Recommendation/commits/main/?author=ShineBo",
    },
    {
        "title": "Portfolio Website",
        "meta": "Owner | Next.js, React, TypeScript, Tailwind CSS, Framer Motion, Vercel | 2026",
        "text": "Designed and deployed a responsive, data-driven portfolio with accessible interactions, motion preferences, theming, metadata, and documented project evidence.",
        "url": "https://portfolio-shinebobo.vercel.app",
    },
    {
        "title": "Flower Classification Research Prototype",
        "meta": "Academic research | Python, TensorFlow, MobileNetV2, Streamlit, OpenCV | 2025",
        "text": "Built a five-class transfer-learning workflow with preprocessing, saved model checkpoints, evaluation visualizations, and a small inference interface.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/lap-flower-classifier",
    },
    {
        "title": "Smart HVAC IoT Controller",
        "meta": "Coursework | C++, ESP32, PlatformIO, Wokwi, MQTT, DHT22 | 2025",
        "text": "Implemented COOL, HEAT, and IDLE states, automatic/manual modes, temperature hysteresis, and MQTT publishing in a repeatable ESP32 simulation.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/iot-project",
    },
    {
        "title": "LAP Portfolio Builder Platform",
        "meta": "Backend contributor | Node.js, Express, MongoDB, Mongoose, JWT, Postman | 2025",
        "text": "Contributed account, profile, blog, comment, and project APIs plus endpoint testing to a collaborative class portfolio platform.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/lap-portfolio-builder-platform",
    },
    {
        "title": "Inventory Management System",
        "meta": "Coursework | Laravel, PHP, MySQL, Blade, Eloquent | 2024",
        "text": "Built a coursework system for products, suppliers, inventory, customers, orders, purchases, authentication, and role-aware workflows.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/inventory-management-system",
    },
    {
        "title": "Hotel Administration System",
        "meta": "Coursework | Java, Spring Boot, Spring Data JPA, Thymeleaf, MySQL | 2025",
        "text": "Created a small hotel-administration prototype for administrator access, room/status management, availability, and check-in records.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/hotel-administration-system",
    },
    {
        "title": "Mini Cafe POS System",
        "meta": "Team coursework | PHP, MySQL, JavaScript, Bootstrap | 2024",
        "text": "Contributed to menu administration, ordering, calculated totals, receipts, and order-history workflows in a group web-programming project.",
        "url": "https://portfolio-shinebobo.vercel.app/projects/mini-cafe-pos-system",
    },
]

ADDITIONAL_PROJECTS = [
    "Country Search - Next.js, React, TypeScript, Material UI, Zustand, and REST Countries API; search, region filters, details, dark mode, and typed client-state practice.",
    "Noodle Town - React, Vite, Tailwind CSS, Framer Motion, and React Router; responsive restaurant menu, reservation, contact, and location interfaces.",
    "Movie Browser - React, React Router, JavaScript, CSS, and TMDB API; movie search, genres, details, and related recommendations.",
    "Weather Forecast App - React, Vite, Tailwind CSS, JavaScript, and OpenWeather API; current conditions and five-day forecast presentation.",
    "Login and Registration Prototype - Next.js, Express, MongoDB, Mongoose, JWT, and bcryptjs; introductory client/server authentication practice.",
    "Nike Landing Page Study - React, Vite, Tailwind CSS, and JavaScript; component composition, responsive spacing, and visual-reference implementation.",
    "Java Flappy Bird - Java, Swing, and AWT; timer-driven movement, keyboard input, collision detection, scoring, and restart behaviour.",
    "Java Hotel Booking GUI - Java, Swing, AWT, and object-oriented programming; room, availability, booking, search, and date-validation workflows.",
]

LEARNING_ARCHIVE = (
    "Additional learning archive includes AI class notebooks, NestJS and Next.js class apps, Streamlit labs, "
    "Laravel and PHP CRUD exercises, Express/MongoDB practice, JavaScript games, and frontend challenge work. "
    "The portfolio documents 35 projects and learning exercises with context and repository evidence where available."
)

ACTIVITIES = [
    "PSU Delegate, Asian Undergraduate Symposium 2026 - NUS College, Singapore, 6-18 Jul 2026; represented PSU's College of Computing in the Communities in Action regional program.",
    "JASSO-Supported Participant, Asia Design Global Workshop 2025 - Shibaura Institute of Technology, Tokyo, 31 Jul-7 Aug 2025; collaborated through ideation, field research, proposal/model development, and final presentation work.",
    "Participant, cross-cultural Data Analysis and Visualization Workshops at PSU Phuket with SIT, Sojo University, Osaka Metropolitan University, and Shunan University in 2024 and 2025.",
    "Supporting Actor and Video Editor, One Billion Rising gender-equality project at Myanmar Institute of Theology.",
    "Guest Speaker, AIESEC in Myanmar Youth and Mental Health podcast for World Mental Health Day 2023.",
    "Youth community volunteer supporting local youth-led activities and informal community initiatives.",
    "Completed 230 recorded PSU activity hours, compared with the 100-hour undergraduate graduation requirement.",
]

DEVELOPMENT_FOCUS = [
    "Full-stack architecture and system-design fundamentals",
    "Networks, cybersecurity, and cryptography foundations",
    "Cloud, deployment, monitoring, and configuration workflows",
    "Applied AI evaluation and responsible model integration",
    "Code quality, testing habits, and technical documentation",
    "Hardware diagnostics and connected hardware-software systems",
]


def set_docx_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_docx_hyperlink(paragraph, text, url, color=BLUE):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Arial")
    run_fonts.set(qn("w:hAnsi"), "Arial")
    properties.append(run_fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    properties.append(size)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    properties.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    label = OxmlElement("w:t")
    label.text = text
    run.append(label)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    paragraph.add_run("Shine Bo Bo | Curriculum Vitae | Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    for element in (begin, instruction, separate, text, end):
        run.append(element)
    paragraph._p.append(run)
    for existing_run in paragraph.runs:
        set_docx_font(existing_run, size=7.5, color=MUTED)


def configure_docx(document):
    section = document.sections[0]
    # Named international_cv_a4 override to compact_reference_guide.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.25)
    normal.font.color.rgb = RGBColor.from_string(SLATE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.06

    heading_tokens = {
        "Heading 1": (14.5, NAVY, 8, 4),
        "Heading 2": (11.2, BLUE, 5, 2),
        "Heading 3": (9.6, NAVY, 3, 1),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_style = styles["List Bullet"]
    bullet_style.font.name = "Arial"
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullet_style.font.size = Pt(9.1)
    bullet_style.font.color.rgb = RGBColor.from_string(SLATE)
    bullet_style.paragraph_format.left_indent = Mm(5.5)
    bullet_style.paragraph_format.first_line_indent = Mm(-3.2)
    bullet_style.paragraph_format.space_after = Pt(1.7)
    bullet_style.paragraph_format.line_spacing = 1.05

    if "CV Meta" not in [style.name for style in styles]:
        meta_style = styles.add_style("CV Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta_style = styles["CV Meta"]
    meta_style.font.name = "Arial"
    meta_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    meta_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    meta_style.font.size = Pt(8.3)
    meta_style.font.color.rgb = RGBColor.from_string(MUTED)
    meta_style.paragraph_format.space_after = Pt(2)
    meta_style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def docx_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_docx_font(run, size=9.1, color=SLATE)


def docx_entry(document, title, meta, bullets, url=None):
    heading = document.add_paragraph(style="Heading 3")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(title)
    set_docx_font(run, size=9.6, color=NAVY, bold=True)
    meta_paragraph = document.add_paragraph(style="CV Meta")
    meta_run = meta_paragraph.add_run(meta)
    set_docx_font(meta_run, size=8.3, color=MUTED, italic=True)
    if url:
        meta_paragraph.add_run(" | ")
        add_docx_hyperlink(meta_paragraph, "View evidence", url)
    for bullet in bullets:
        docx_bullet(document, bullet)


def build_docx():
    DOCX_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_docx(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("SHINE BO BO")
    set_docx_font(title_run, size=25, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(3)
    subtitle_run = subtitle.add_run("CURRICULUM VITAE | SOFTWARE AND DIGITAL ENGINEERING")
    set_docx_font(subtitle_run, size=10, color=BLUE, bold=True)

    contact = document.add_paragraph()
    contact.paragraph_format.space_after = Pt(1)
    contact.add_run("Phuket, Thailand | +66 61 858 1574 | ")
    add_docx_hyperlink(contact, "shinebobo648@gmail.com", "mailto:shinebobo648@gmail.com")
    links = document.add_paragraph()
    links.paragraph_format.space_after = Pt(5)
    add_docx_hyperlink(links, "Portfolio", "https://portfolio-shinebobo.vercel.app")
    links.add_run(" | ")
    add_docx_hyperlink(links, "GitHub", "https://github.com/ShineBo")
    links.add_run(" | ")
    add_docx_hyperlink(links, "LinkedIn", "https://www.linkedin.com/in/s02bb/")
    for paragraph in (contact, links):
        for run in paragraph.runs:
            set_docx_font(run, size=8.5, color=SLATE)

    document.add_heading("Professional Profile", level=1)
    document.add_paragraph(PROFILE)

    document.add_heading("Education", level=1)
    for item in EDUCATION:
        docx_entry(document, item["title"], item["meta"], item["bullets"])

    document.add_heading("Technical Competencies", level=1)
    note = document.add_paragraph()
    note_run = note.add_run(
        "These skills reflect project, coursework, research, or hands-on work exposure - not claims of professional mastery."
    )
    set_docx_font(note_run, size=8.4, color=MUTED, italic=True)
    for label, details in SKILLS:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_docx_font(label_run, size=9.1, color=NAVY, bold=True)
        details_run = paragraph.add_run(details)
        set_docx_font(details_run, size=9.1, color=SLATE)

    document.add_heading("Experience and Leadership", level=1)
    for item in EXPERIENCE[:2]:
        docx_entry(document, item["title"], item["meta"], item["bullets"])

    document.add_page_break()

    document.add_heading("Experience and Leadership - Continued", level=1)
    for item in EXPERIENCE[2:]:
        docx_entry(document, item["title"], item["meta"], item["bullets"])

    document.add_heading("Selected Technical Projects", level=1)
    for item in PROJECTS[:7]:
        docx_entry(document, item["title"], item["meta"], [item["text"]], item["url"])

    document.add_page_break()

    document.add_heading("Selected Technical Projects - Continued", level=1)
    for item in PROJECTS[7:]:
        docx_entry(document, item["title"], item["meta"], [item["text"]], item["url"])

    document.add_heading("Additional Project Work", level=1)
    for item in ADDITIONAL_PROJECTS:
        docx_bullet(document, item)
    archive_note = document.add_paragraph(LEARNING_ARCHIVE)
    for run in archive_note.runs:
        set_docx_font(run, size=8.5, color=MUTED, italic=True)

    document.add_heading("International and Community Engagement", level=1)
    for item in ACTIVITIES:
        docx_bullet(document, item)

    document.add_heading("Languages", level=1)
    document.add_paragraph("Burmese - Native | English - Fluent | Korean - Basic coursework")

    document.add_heading("Current Development Focus", level=1)
    for item in DEVELOPMENT_FOCUS:
        docx_bullet(document, item)

    document.core_properties.title = "Shine Bo Bo - Curriculum Vitae"
    document.core_properties.subject = "Software and Digital Engineering curriculum vitae"
    document.core_properties.author = "Shine Bo Bo"
    document.core_properties.keywords = "Software Development, Digital Engineering, Computer Science, Internship"
    document.save(DOCX_OUTPUT)
    return DOCX_OUTPUT


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "CV Name",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=25,
            textColor=colors.HexColor(f"#{NAVY}"),
            alignment=TA_LEFT,
            spaceAfter=1,
        ),
        "subtitle": ParagraphStyle(
            "CV Subtitle",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.6,
            leading=12,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "CV Contact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.1,
            leading=10.5,
            textColor=colors.HexColor(f"#{SLATE}"),
        ),
        "section": ParagraphStyle(
            "CV Section",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=11.3,
            leading=13,
            textColor=colors.HexColor(f"#{NAVY}"),
            spaceBefore=5,
            spaceAfter=2,
        ),
        "entry": ParagraphStyle(
            "CV Entry",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9.1,
            leading=11,
            textColor=colors.HexColor(f"#{NAVY}"),
            spaceAfter=1,
        ),
        "meta": ParagraphStyle(
            "CV Meta PDF",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=7.8,
            leading=9.5,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=1,
        ),
        "body": ParagraphStyle(
            "CV Body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.25,
            leading=10.2,
            textColor=colors.HexColor(f"#{SLATE}"),
            alignment=TA_LEFT,
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "CV Bullet",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.1,
            leading=10,
            leftIndent=8,
            firstLineIndent=-7,
            bulletIndent=0,
            textColor=colors.HexColor(f"#{SLATE}"),
            spaceAfter=1.1,
        ),
        "small": ParagraphStyle(
            "CV Small",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=7.7,
            leading=9.4,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=2,
        ),
    }


PDF_STYLES = pdf_styles()


def pdf_section(story, title):
    story.append(Paragraph(escape(title.upper()), PDF_STYLES["section"]))
    story.append(
        HRFlowable(
            width="100%",
            thickness=0.65,
            color=colors.HexColor(f"#{BLUE}"),
            spaceBefore=0,
            spaceAfter=3,
        )
    )


def pdf_entry(title, meta, bullets, url=None):
    meta_text = escape(meta)
    if url:
        meta_text += f" | <link href='{escape(url)}' color='#{BLUE}'>View evidence</link>"
    blocks = [
        Paragraph(escape(title), PDF_STYLES["entry"]),
        Paragraph(meta_text, PDF_STYLES["meta"]),
    ]
    blocks.extend(Paragraph(escape(item), PDF_STYLES["bullet"], bulletText="-") for item in bullets)
    blocks.append(Spacer(1, 1.5))
    return KeepTogether(blocks)


def pdf_page_frame(canvas, document):
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(colors.HexColor(f"#{BLUE}"))
    canvas.setLineWidth(1.1)
    canvas.line(17 * mm, height - 10 * mm, width - 17 * mm, height - 10 * mm)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.setFont("Helvetica", 7.1)
    canvas.drawString(17 * mm, 8.5 * mm, "Shine Bo Bo | Curriculum Vitae")
    canvas.drawRightString(width - 17 * mm, 8.5 * mm, f"Page {document.page}")
    canvas.restoreState()


def build_pdf():
    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    PUBLIC_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=17 * mm,
        leftMargin=17 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title="Shine Bo Bo - Curriculum Vitae",
        author="Shine Bo Bo",
        subject="Software and Digital Engineering curriculum vitae",
        creator="Shine Bo Bo portfolio",
    )

    story = [
        Paragraph("SHINE BO BO", PDF_STYLES["name"]),
        Paragraph("CURRICULUM VITAE | SOFTWARE AND DIGITAL ENGINEERING", PDF_STYLES["subtitle"]),
        Paragraph(
            "Phuket, Thailand | +66 61 858 1574 | "
            "<link href='mailto:shinebobo648@gmail.com' color='#2563EB'>shinebobo648@gmail.com</link>",
            PDF_STYLES["contact"],
        ),
        Paragraph(
            "<link href='https://portfolio-shinebobo.vercel.app' color='#2563EB'>Portfolio</link> | "
            "<link href='https://github.com/ShineBo' color='#2563EB'>GitHub</link> | "
            "<link href='https://www.linkedin.com/in/s02bb/' color='#2563EB'>LinkedIn</link>",
            PDF_STYLES["contact"],
        ),
        Spacer(1, 4),
    ]

    pdf_section(story, "Professional Profile")
    story.append(Paragraph(escape(PROFILE), PDF_STYLES["body"]))

    pdf_section(story, "Education")
    for item in EDUCATION:
        story.append(pdf_entry(item["title"], item["meta"], item["bullets"]))

    pdf_section(story, "Technical Competencies")
    story.append(
        Paragraph(
            "Skills reflect project, coursework, research, or hands-on work exposure - not claims of professional mastery.",
            PDF_STYLES["small"],
        )
    )
    for label, details in SKILLS:
        story.append(
            Paragraph(
                f"<b>{escape(label)}:</b> {escape(details)}",
                PDF_STYLES["body"],
            )
        )

    pdf_section(story, "Experience and Leadership")
    for item in EXPERIENCE[:2]:
        story.append(pdf_entry(item["title"], item["meta"], item["bullets"]))

    story.append(PageBreak())
    pdf_section(story, "Experience and Leadership - Continued")
    for item in EXPERIENCE[2:]:
        story.append(pdf_entry(item["title"], item["meta"], item["bullets"]))

    pdf_section(story, "Selected Technical Projects")
    for item in PROJECTS[:7]:
        story.append(pdf_entry(item["title"], item["meta"], [item["text"]], item["url"]))

    story.append(PageBreak())
    pdf_section(story, "Selected Technical Projects - Continued")
    for item in PROJECTS[7:]:
        story.append(pdf_entry(item["title"], item["meta"], [item["text"]], item["url"]))

    pdf_section(story, "Additional Project Work")
    for item in ADDITIONAL_PROJECTS:
        story.append(Paragraph(escape(item), PDF_STYLES["bullet"], bulletText="-"))
    story.append(Paragraph(escape(LEARNING_ARCHIVE), PDF_STYLES["small"]))

    pdf_section(story, "International and Community Engagement")
    for item in ACTIVITIES:
        story.append(Paragraph(escape(item), PDF_STYLES["bullet"], bulletText="-"))

    pdf_section(story, "Languages")
    story.append(Paragraph("Burmese - Native | English - Fluent | Korean - Basic coursework", PDF_STYLES["body"]))

    pdf_section(story, "Current Development Focus")
    for item in DEVELOPMENT_FOCUS:
        story.append(Paragraph(escape(item), PDF_STYLES["bullet"], bulletText="-"))

    document.build(story, onFirstPage=pdf_page_frame, onLaterPages=pdf_page_frame)
    copy2(PDF_OUTPUT, PUBLIC_OUTPUT)
    return PDF_OUTPUT


def main():
    docx_path = build_docx()
    pdf_path = build_pdf()
    print(f"Generated {docx_path}")
    print(f"Generated {pdf_path}")
    print(f"Published {PUBLIC_OUTPUT}")


if __name__ == "__main__":
    main()
